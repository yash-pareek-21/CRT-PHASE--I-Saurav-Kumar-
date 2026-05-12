from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.platypus.flowables import HRFlowable

from docx import Document
from docx.shared import Pt

# Resume Data
resume_text = {
    "name": "YASH PAREEK",
    "title": "Frontend Developer Intern Applicant",

    "contact": """
Jaipur, Rajasthan
Email: youremail@gmail.com
Phone: +91-XXXXXXXXXX
GitHub: github.com/yourprofile
LinkedIn: linkedin.com/in/yourprofile
""",

    "objective": """
Passionate and detail-oriented B.Tech Computer Science student at JECRC University currently in 6th semester, seeking a Frontend Developer Internship. Skilled in HTML, CSS, JavaScript, and basic React.js with a strong interest in building responsive and user-friendly web applications.
""",

    "education": """
Bachelor of Technology (B.Tech) – Computer Science Engineering
JECRC University, Jaipur
2023 – 2027 (Expected)
""",

    "skills": [
        "HTML5",
        "CSS3",
        "JavaScript",
        "React.js (Basics)",
        "C++",
        "Git & GitHub",
        "Responsive Web Design",
        "Problem Solving",
        "Team Collaboration"
    ],

    "projects": [
        "Portfolio Website – Built a responsive portfolio using HTML, CSS, and JavaScript.",
        "College Management System – Developed using C++ with inheritance and virtual functions.",
        "To-Do List Web App – Created a task management app with add/delete functionality."
    ],

    "achievements": [
        "Solved coding problems in C++ and DSA.",
        "Participated in technical workshops and learning programs.",
        "Active learner in frontend development."
    ],

    "interests": [
        "Web Development",
        "UI/UX Design",
        "Open Source Learning"
    ]
}

# File Names
docx_path = "Yash_Pareek_Resume.docx"
pdf_path = "Yash_Pareek_Resume.pdf"

# ---------------- DOCX CREATION ---------------- #

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading(resume_text["name"], level=1)
doc.add_paragraph(resume_text["title"])
doc.add_paragraph(resume_text["contact"])

# Career Objective
doc.add_heading("Career Objective", level=2)
doc.add_paragraph(resume_text["objective"])

# Education
doc.add_heading("Education", level=2)
doc.add_paragraph(resume_text["education"])

# Skills
doc.add_heading("Skills", level=2)
for skill in resume_text["skills"]:
    doc.add_paragraph(skill, style='List Bullet')

# Projects
doc.add_heading("Projects", level=2)
for project in resume_text["projects"]:
    doc.add_paragraph(project, style='List Bullet')

# Achievements
doc.add_heading("Achievements / Certifications", level=2)
for achievement in resume_text["achievements"]:
    doc.add_paragraph(achievement, style='List Bullet')

# Interests
doc.add_heading("Extra-Curricular / Interests", level=2)
for interest in resume_text["interests"]:
    doc.add_paragraph(interest, style='List Bullet')

# Save DOCX
doc.save(docx_path)

# ---------------- PDF CREATION ---------------- #

doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter)

styles = getSampleStyleSheet()

elements = []

# Title
elements.append(Paragraph(f"<b>{resume_text['name']}</b>", styles['Title']))
elements.append(Paragraph(resume_text["title"], styles['Normal']))
elements.append(Paragraph(resume_text["contact"], styles['Normal']))
elements.append(Spacer(1, 12))


# Function to add sections
def add_section(title, content):
    elements.append(HRFlowable(width="100%"))
    elements.append(Paragraph(f"<b>{title}</b>", styles['Heading2']))

    if isinstance(content, list):
        for item in content:
            elements.append(Paragraph(f"• {item}", styles['BodyText']))
    else:
        elements.append(Paragraph(content, styles['BodyText']))

    elements.append(Spacer(1, 10))


# Add Sections
add_section("Career Objective", resume_text["objective"])
add_section("Education", resume_text["education"])
add_section("Skills", resume_text["skills"])
add_section("Projects", resume_text["projects"])
add_section("Achievements / Certifications", resume_text["achievements"])
add_section("Extra-Curricular / Interests", resume_text["interests"])

# Build PDF
doc_pdf.build(elements)

print("Resume created successfully!")
print("DOCX File:", docx_path)
print("PDF File:", pdf_path)